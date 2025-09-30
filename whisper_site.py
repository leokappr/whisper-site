import os
import json
import streamlit as st
import openai
from pathlib import Path

# exports
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm

# pour le bouton copier (JS)
import streamlit.components.v1 as components

openai.api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Whisper Transcripteur", layout="wide")
st.title("🎧 Whisper Transcripteur")

uploaded_file = st.file_uploader("📁 Importer un fichier audio", type=["mp3", "m4a", "wav"])

def make_paragraphs(text: str):
    """
    Retourne une liste de paragraphes propres :
    – normalise les retours à la ligne
    – supprime les lignes vides multiples
    """
    # remplace les retours Windows
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    # si pas de doubles lignes vides, on coupe à chaque ligne ; sinon on garde les blocs
    if "\n\n" in t:
        parts = [p.strip() for p in t.split("\n\n") if p.strip()]
    else:
        parts = [p.strip() for p in t.split("\n") if p.strip()]
    return parts

def create_docx(content: str, out_path: Path):
    doc = Document()
    doc.add_heading("Transcription", level=1)
    for p in make_paragraphs(content):
        doc.add_paragraph(p)
    doc.save(out_path)

def create_pdf(content: str, out_path: Path):
    # PDF propre : marges, interligne, paragraphes
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        'NormalPlus',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=12,
        leading=18,        # interligne
        spaceAfter=10      # espace après chaque par.
    )
    h1 = styles['Heading1']

    story = [Paragraph("Transcription", h1), Spacer(1, 12)]
    for p in make_paragraphs(content):
        # échappe minimalement les chevrons
        p_html = p.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(p_html, normal))
        story.append(Spacer(1, 6))

    doc.build(story)

if uploaded_file:
    with st.spinner("🧠 Transcription en cours..."):
        # sauvegarde temporaire
        tmp_audio = Path(uploaded_file.name)
        with open(tmp_audio, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # appel Whisper (transcription brute)
        with open(tmp_audio, "rb") as af:
            transcript = openai.audio.transcriptions.create(
                model="whisper-1",
                file=af,
                response_format="text"
            )

    st.success("✅ Transcription terminée !")

    # 📝 zone éditable (tu peux corriger avant export)
    edited_text = st.text_area("📝 Transcription (modifiable avant téléchargement)", transcript, height=420)

    # 📋 bouton COPIER (presse-papiers)
    col_copy, col_sp, col_docx, col_pdf = st.columns([1, 0.2, 1, 1])
    with col_copy:
        if st.button
import streamlit as st
import openai
import os
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm

openai.api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Whisper Transcripteur", layout="wide")

st.title("🎧 Whisper Transcripteur")

uploaded_file = st.file_uploader("📁 Importer un fichier audio", type=["mp3", "m4a", "wav"])

if uploaded_file:
    with st.spinner("🧠 Transcription en cours..."):
        # Enregistre temporairement le fichier audio
        with open(uploaded_file.name, "wb") as f:
            f.write(uploaded_file.getbuffer())

        audio_file = open(uploaded_file.name, "rb")

        # Transcription avec Whisper
        transcript = openai.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            response_format="text"
        )

        st.success("✅ Transcription terminée !")

        # 🧩 Zone d’édition : l’utilisateur peut relire et modifier
        edited_text = st.text_area("📝 Votre transcription (modifiable avant téléchargement)", transcript, height=400)

        # 🧾 Téléchargement en DOCX
        def create_docx(content):
            doc = Document()
            doc.add_heading("Transcription", level=1)
            for paragraph in content.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.save("transcription.docx")
            return "transcription.docx"

        # 🧾 Téléchargement en PDF bien formaté
        def create_pdf(content):
            pdf_path = "transcription.pdf"
            doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                                    rightMargin=2*cm, leftMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            normal = ParagraphStyle(
                'Normal',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=12,
                leading=18,
                spaceAfter=12
            )

            story = [Paragraph("<b>Transcription</b>", styles['Heading1']), Spacer(1, 12)]
            for paragraph in content.split("\n"):
                if paragraph.strip():
                    story.append(Paragraph(paragraph.strip(), normal))
                    story.append(Spacer(1, 8))

            doc.build(story)
            return pdf_path

        # 🧱 Boutons de téléchargement
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📄 Télécharger en DOCX"):
                docx_file = create_docx(edited_text)
                with open(docx_file, "rb") as f:
                    st.download_button("⬇️ Télécharger DOCX", f, file_name="transcription.docx")

        with col2:
            if st.button("🧾 Télécharger en PDF"):
                pdf_file = create_pdf(edited_text)
                with open(pdf_file, "rb") as f:
                    st.download_button("⬇️ Télécharger PDF", f, file_name="transcription.pdf")
import os
import io
import json
import math
from datetime import datetime
from pathlib import Path
from openai import OpenAI
import streamlit as st
from pydub import AudioSegment
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

# 🔑 Récupération clé API
api_key = os.getenv("OPENAI_API_KEY") or os.getenv("WISPER")
if not api_key:
    st.error("❌ Clé API non trouvée. Vérifie ton fichier ~/.zshrc.")
    st.stop()

client = OpenAI(api_key=api_key)

# 📁 Dossiers
OUTPUT_DIR = Path.home() / "Whisper_Transcriptions"
OUTPUT_DIR.mkdir(exist_ok=True)
HISTORIQUE_FILE = OUTPUT_DIR / "historique.json"

# 🔄 Charger historique existant
def charger_historique():
    if HISTORIQUE_FILE.exists():
        with open(HISTORIQUE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

# 💾 Sauvegarder historique
def sauvegarder_historique(data):
    with open(HISTORIQUE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ✂️ Fonction pour découper les fichiers audio > 25 Mo
def decouper_audio(path):
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb <= 25:
        return [path]
    
    st.warning(f"⚠️ Fichier trop volumineux ({size_mb:.1f} Mo) → découpage en segments...")
    audio = AudioSegment.from_file(path)
    duree_ms = len(audio)
    nb_segments = math.ceil(size_mb / 25)
    segment_duree = duree_ms // nb_segments
    segments = []

    for i in range(nb_segments):
        start = i * segment_duree
        end = (i + 1) * segment_duree if i < nb_segments - 1 else duree_ms
        segment = audio[start:end]
        segment_path = path.parent / f"{path.stem}_part{i+1}.mp3"
        segment.export(segment_path, format="mp3")
        segments.append(segment_path)

    return segments

# 🧾 Génération du .docx
def save_docx(text, output_path):
    doc = Document()
    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.save(output_path)

# 🧾 Génération du .pdf
def save_pdf(text, output_path):
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 50
    for line in text.split("\n"):
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line)
        y -= 15
    c.save()

# ⚙️ Transcription principale
def transcrire_audio(path):
    segments = decouper_audio(path)
    texte_total = ""
    for i, seg in enumerate(segments, 1):
        with open(seg, "rb") as f:
            st.info(f"🧠 Transcription du segment {i}/{len(segments)}...")
            res = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="text",
                language="fr"
            )
            texte_total += res + "\n"

    return texte_total.strip()

# 🎨 CONFIG PAGE
st.set_page_config(page_title="Transcripteur Whisper 🎧", layout="wide")
menu = st.sidebar.radio("📋 Menu", ["🎙️ Transcrire", "📚 Historique", "⚙️ Paramètres"])

if menu == "🎙️ Transcrire":
    st.title("🎧 Transcripteur Audio Whisper")
    st.write("Upload un fichier audio (mp3, m4a, wav). Fichiers longs acceptés (1h max).")

    file = st.file_uploader("📂 Choisis ton fichier", type=["mp3", "m4a", "wav"])
    if file:
        temp_path = Path(file.name)
        with open(temp_path, "wb") as f:
            f.write(file.getbuffer())
        st.success(f"✅ Fichier chargé : {file.name}")

        if st.button("🎙️ Lancer la transcription"):
            with st.spinner("Transcription en cours..."):
                texte = transcrire_audio(temp_path)

                # Enregistrement
                base_name = temp_path.stem + "_" + datetime.now().strftime("%Y%m%d_%H%M%S")
                txt_path = OUTPUT_DIR / f"{base_name}.txt"
                docx_path = OUTPUT_DIR / f"{base_name}.docx"
                pdf_path = OUTPUT_DIR / f"{base_name}.pdf"

                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(texte)
                save_docx(texte, docx_path)
                save_pdf(texte, pdf_path)

                st.success("✅ Transcription terminée !")
                st.text_area("📝 Résultat :", texte, height=300)
                st.download_button("⬇️ Télécharger (.docx)", data=docx_path.read_bytes(), file_name=docx_path.name)
                st.download_button("⬇️ Télécharger (.pdf)", data=pdf_path.read_bytes(), file_name=pdf_path.name)
                st.code(str(txt_path))

                # Historique
                histo = charger_historique()
                histo.append({
                    "nom": file.name,
                    "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "fichier_txt": str(txt_path)
                })
                sauvegarder_historique(histo)

elif menu == "📚 Historique":
    st.title("📚 Historique des transcriptions")
    histo = charger_historique()
    if not histo:
        st.info("Aucune transcription enregistrée pour le moment.")
    else:
        for item in reversed(histo):
            with st.expander(f"📝 {item['nom']} ({item['date']})"):
                path = Path(item["fichier_txt"])
                if path.exists():
                    texte = path.read_text(encoding="utf-8")
                    st.text_area("Texte :", texte, height=200)
                    st.download_button("⬇️ Télécharger (.txt)", data=path.read_bytes(), file_name=path.name)
                else:
                    st.warning("Fichier manquant sur le disque.")

elif menu == "⚙️ Paramètres":
    st.title("⚙️ Paramètres")
    st.write(f"📂 Dossier de sauvegarde : `{OUTPUT_DIR}`")
    st.write("🧠 Modèle : Whisper-1")
    st.write("🌍 Langue : Auto")
    st.write("💾 Historique enregistré localement")
    st.success("✅ Configuration prête !")

